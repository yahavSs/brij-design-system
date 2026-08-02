"""
brij_doc — Brand-styled .docx generator for Brij.

Usage from a Claude-written script:

    from brij_doc import BrijDocument

    doc = BrijDocument(title="Brij Signal One-Pager")
    doc.eyebrow("01 — INTRODUCTION")
    doc.display("Time to brij.")
    doc.body("Sharp, direct copy goes here.")
    doc.headline("What we do")
    doc.bullet_list(["Diagnose.", "Build.", "Hand off."])
    doc.callout("We are not a vendor. We are a home lab.")
    doc.divider()
    doc.save("output.docx")

All styles map to the Brij brand: Plus Jakarta Sans + Arbutus Slab,
ink text, sentence case, brackets for eyebrows, no shadows or emoji.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ───── Brand tokens ─────
INK = RGBColor(0x18, 0x1B, 0x18)
T2 = RGBColor(0x6B, 0x6B, 0x63)
T3 = RGBColor(0xA3, 0xA3, 0x9B)
ORANGE = RGBColor(0xFF, 0xCE, 0x65)
CREAM_HEX = "FAF9F4"
SAGE_HEX = "EBEDE7"
BORDER = RGBColor(0xE8, 0xE6, 0xE1)

SANS = "Plus Jakarta Sans"
SERIF = "Arbutus Slab"


def _set_run(run, *, font=SANS, size=11, color=INK, bold=False, italic=False, tracking=0):
    """Apply brand-aware formatting to a run."""
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if tracking:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(tracking))
        rpr.append(spacing)


def _set_para(p, *, space_before=0, space_after=8, line_spacing=1.4, alignment=None, keep_with_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    if keep_with_next:
        pf.keep_with_next = True


class BrijDocument:
    """A Brij-branded Word document. Methods build sections in order."""

    def __init__(self, title=None, page_size="A4"):
        self.doc = Document()
        self._configure_page(page_size)
        self._set_cream_background()
        if title:
            self.doc.core_properties.title = title

    def _configure_page(self, size):
        section = self.doc.sections[0]
        if size.upper() == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        else:  # Letter
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    def _set_cream_background(self):
        """Apply cream page background (visible in Word, ignored by Google Docs)."""
        bg = OxmlElement("w:background")
        bg.set(qn("w:color"), CREAM_HEX)
        self.doc.element.insert(0, bg)
        display_bg = OxmlElement("w:displayBackgroundShape")
        settings = self.doc.settings.element
        settings.append(display_bg)

    # ───── Content blocks ─────

    def eyebrow(self, text):
        """Small uppercase tracked label, e.g. '[ 01 — FOUNDATION ]'."""
        p = self.doc.add_paragraph()
        run = p.add_run(text.upper())
        _set_run(run, size=9, color=T3, tracking=24)
        _set_para(p, space_before=24, space_after=12, keep_with_next=True)
        return self

    def display(self, text):
        """The biggest title. Plus Jakarta Sans, ink, bold."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        _set_run(run, size=32, color=INK, bold=True)
        _set_para(p, space_before=8, space_after=18, line_spacing=1.1, keep_with_next=True)
        return self

    def headline(self, text):
        """Section headline. Plus Jakarta Sans 22pt."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        _set_run(run, size=22, color=INK, bold=True)
        _set_para(p, space_before=24, space_after=12, line_spacing=1.2, keep_with_next=True)
        return self

    def title(self, text):
        """Sub-section title. Plus Jakarta Sans 14pt."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        _set_run(run, size=14, color=INK, bold=True)
        _set_para(p, space_before=14, space_after=6, keep_with_next=True)
        return self

    def serif_title(self, text):
        """Accent title in Arbutus Slab. For panel titles, callouts."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        _set_run(run, font=SERIF, size=18, color=INK)
        _set_para(p, space_before=18, space_after=8, keep_with_next=True)
        return self

    def body(self, text, color=None):
        """Default paragraph. 11pt body."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        _set_run(run, size=11, color=color or INK)
        _set_para(p, space_after=10)
        return self

    def lead(self, text):
        """Larger intro paragraph. Used right after a display."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        _set_run(run, size=13, color=T2)
        _set_para(p, space_after=14, line_spacing=1.5)
        return self

    def caption(self, text):
        """Small muted caption."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        _set_run(run, size=9, color=T3)
        _set_para(p, space_after=8)
        return self

    def bullet_list(self, items):
        """Sentence-case bullets, ink text."""
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            _set_run(run, size=11, color=INK)
            _set_para(p, space_after=4)
        return self

    def numbered_list(self, items):
        for item in items:
            p = self.doc.add_paragraph(style="List Number")
            run = p.add_run(item)
            _set_run(run, size=11, color=INK)
            _set_para(p, space_after=4)
        return self

    def callout(self, text):
        """A cream-bg pull-quote-like block. Arbutus Slab, larger."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.rows[0].cells[0]
        # Cell shading: sage at low opacity (we approximate with a light tint).
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F1EB")  # mid between cream and sage
        tc_pr.append(shd)
        # No table border
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tc_pr.append(border)
        # Padding
        tcMar = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), "320")
            m.set(qn("w:type"), "dxa")
            tcMar.append(m)
        tc_pr.append(tcMar)
        # Content
        p = cell.paragraphs[0]
        run = p.add_run(text)
        _set_run(run, font=SERIF, size=16, color=INK)
        _set_para(p, space_before=4, space_after=4, line_spacing=1.35)
        # Spacer after
        spacer = self.doc.add_paragraph()
        _set_para(spacer, space_after=8)
        return self

    def panel(self, lines):
        """A subtle panel (sage-tint cell with multiple lines).

        lines: list of (kind, text) tuples where kind is 'title' | 'body' | 'caption'.
        """
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4EE")
        tc_pr.append(shd)
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tc_pr.append(border)
        tcMar = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), "260")
            m.set(qn("w:type"), "dxa")
            tcMar.append(m)
        tc_pr.append(tcMar)
        # Replace default first paragraph
        first = True
        for kind, text in lines:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            run = p.add_run(text)
            if kind == "title":
                _set_run(run, size=12, bold=True, color=INK)
                _set_para(p, space_after=4)
            elif kind == "caption":
                _set_run(run, size=9, color=T3)
                _set_para(p, space_after=2)
            else:
                _set_run(run, size=10, color=INK)
                _set_para(p, space_after=4)
        # Spacer
        spacer = self.doc.add_paragraph()
        _set_para(spacer, space_after=6)
        return self

    def divider(self):
        """A thin border line for visual separation."""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "E8E6E1")
        pBdr.append(bottom)
        pPr.append(pBdr)
        _set_para(p, space_before=10, space_after=10)
        return self

    def page_break(self):
        self.doc.add_page_break()
        return self

    def save(self, path):
        self.doc.save(path)
        return path
